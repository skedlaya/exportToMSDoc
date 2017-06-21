import sys
import requests
from docx import Document
from docx.shared import RGBColor

#Transaction ID as first input argument
transactionID = sys.argv[1]

#API Key as second input argument
apiKey = sys.argv[2]

#url to call the API
url = "https://api.capio.ai/v1/speech/transcript/" + transactionID

# Get request to retrieve JSON from API
r = requests.get(url, headers={"apiKey":apiKey});
data = r.json()

# Create MS docx document
document = Document()

# Convert float value of time in day:hour:minute.second format
def getTime(floatTime):
    roundTime = round(floatTime,3)
    s = (roundTime - int(roundTime)+0.005)*100
    h, m = divmod(int(floatTime), 60)
    d, h = divmod(h, 12)
    return "%02d:%02d:%02d.%02d" % (d, h, m, s)

# If request is successful, save data in MS docx
if r.status_code == 200:
    # Iterate over each word in the JSON
    for d in data:
        for res in d["result"]:
            for alt in res["alternative"]:
                
                #Add a new paragraph for a transcript
                p = document.add_paragraph('')
                
                for i, word in enumerate(alt["words"]):
                    
                    #Retrieve time stamp from the first word
                    if i==0:
                        run = p.add_run(getTime(word["from"]) + "\t")
                        
                        # Set blue color for timestamp
                        font = run.font
                        font.color.rgb = RGBColor(0x00, 0x00, 0xff)

                    #Add word to time stamp
                    run = p.add_run(word["word"]+ " ")
                    font = run.font
        
                    if word["confidence"] < 0.75:
                        # Red color
                        font.color.rgb = RGBColor(0xff, 0x00, 0x00)
                    else:
                        # Black color
                        font.color.rgb = RGBColor(0x00, 0x00, 0x00)

# Save MS word document
document.save('results_created.docx')



