from docx import Document
import os

#Run the program and to create the file results_created.docx
def createDoc():
    os.system('python exporter.py 593f237fbcae700012ba8fcd 262ac9a0c9ba4d179aad4c0b9b02120a')

# Tests time stamp and transcript
# Does not test word confidence
def test_answer():
    createDoc()
    document = Document('results_created.docx')
    document_ref = Document('results_reference.docx')
    for i, para in enumerate(document.paragraphs):
	assert document.paragraphs[i].text.strip().split() == document_ref.paragraphs[i].text.strip().split()

